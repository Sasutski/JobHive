import os
import json
from pathlib import Path
import google.generativeai as genai
import base64
from google.generativeai import types
from rich.console import Console
from rich.panel import Panel
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Assuming the presence of some classes and methods (e.g., storage_manager, console) from your setup.
console = Console()

# Configure Generative AI API
API_KEY = 'AIzaSyBlK25lcxrb9krNnh0PwDFyyae9Vn6rNqA'
client = genai.configure(api_key=API_KEY)

class ResumeReviewer:
    def __init__(self):
        pass  # No need to handle uploading, assuming it's already done
        self.console = Console()  # Initialize console as instance variable
    def open_file_dialog(self):
        """Simulate file selection in a dialog."""
        # This would integrate with your actual file dialog logic
        return "path_to_selected_resume.pdf"

    def review_resume(self, resume_path: str, model_resume_url: str):
        """Review the candidate resume against the model resume."""
        try:
            # Check if the resume is a valid PDF or Word document
            if not resume_path.lower().endswith(('.pdf', '.doc', '.docx')):
                console.print("[red]Unsupported file type. Please use PDF or Word document.[/red]")
                return

            # No need to upload the file, we just proceed to the review step
            console.print(Panel(
                "[cyan]Analyzing your resume against job requirements...\nThis may take a moment.[/cyan]",
                title="Resume Review",
                border_style="blue",
                width=60
            ))

            # Directly process the review for the locally selected candidate resume and model resume URL
            self.analyze_resumes(resume_path, model_resume_url)

        except Exception as e:
            console.print(Panel(
                f"[red]Unexpected error: {str(e)}[/red]\n\n" +
                "[yellow]Please try again later or contact support if the issue persists.[/yellow]",
                title="Error",
                border_style="red",
                width=60
            ))

    def analyze_resumes(self, candidate_resume_path: str, model_resume_url: str):
        """Analyze and compare the candidate's resume against the employer's model resume."""
        try:
            if not os.path.exists(candidate_resume_path):
                console.print("[red]Error: Candidate resume file not found[/red]")
                return

            # Download the model resume from the URL
            model_resume_path = self.download_resume(model_resume_url)
            if not model_resume_path:
                console.print("[red]Error: Could not download model resume[/red]")
                return

            # Read the resumes as bytes
            try:
                with open(candidate_resume_path, "rb") as doc_file:
                    candidate_bytes = doc_file.read()
                with open(model_resume_path, "rb") as model_file:
                    employer_bytes = model_file.read()
            except Exception as e:
                console.print(f"[red]Error reading resume files: {str(e)}[/red]")
                return

            # Encode to base64 for prompt construction
            candidate_data = base64.standard_b64encode(candidate_bytes).decode("utf-8")
            employer_data = base64.standard_b64encode(employer_bytes).decode("utf-8")

            # Build prompt
            prompt_text = (
                "You are a professional resume reviewer. "
                "Please compare the following candidate's resume with the employer's model resume. "
                "Provide a brief summary of the employer's model resume and its strengths. "
                "Then provide detailed feedback on the candidate's resume, highlighting strengths, weaknesses, "
                "and areas for improvement. Be specific and constructive.\n\n"
                f"Employer's Model Resume (base64):\n{employer_data}\n\n"
                f"Candidate's Resume (base64):\n{candidate_data}\n\n"
                "Feedback:"
            )

            # Create Part objects using helper methods
            employer_part = types.Part.from_bytes(data=employer_bytes, mime_type="application/pdf")
            candidate_part = types.Part.from_bytes(data=candidate_bytes, mime_type="application/pdf")
            prompt_part = types.Part.from_text(text=prompt_text)

            # Call the API using generate_content_stream
            model = genai.GenerativeModel('gemini-1.5-flash')
            stream = model.generate_content_stream(
                contents=[employer_part, candidate_part, prompt_part]
            )

            # Collect feedback chunks
            feedback_text = ""
            console.print("\nFeedback:\n")
            for chunk in stream:
                feedback_text += chunk.text
                console.print(chunk.text, end="")
            console.print("\n")  # Add a newline after feedback

            # Save feedback as Word document
            self.save_feedback_as_word(feedback_text, candidate_resume_path)

        except Exception as e:
            console.print(f"[red]Error comparing resumes: {str(e)}[/red]")
        finally:
            # Clean up temporary files
            if model_resume_path and os.path.exists(model_resume_path):
                try:
                    os.remove(model_resume_path)
                except:
                    pass

    def save_feedback_as_word(self, feedback_text: str, resume_path: str):
            """Save the feedback as a formatted Word document with user-selected location."""
            try:
                # Create a new Word document
                doc = Document()
        
                # Add title
                doc.add_heading('Resume Review Feedback', 0)
        
                # Add timestamp
                doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
                # Add a line for the reviewed resume
                doc.add_paragraph(f'Resume reviewed: {Path(resume_path).name}')
        
                # Add horizontal line
                doc.add_paragraph('_' * 50)
        
                # Add feedback content
                doc.add_heading('Detailed Feedback', 1)
                doc.add_paragraph(feedback_text)
        
                # Get original filename
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                original_filename = f'resume_feedback_{timestamp}.docx'
                    
                # Let user choose save location
                self.console.print("\n[yellow]Where would you like to save the feedback document?[/yellow]")
                self.console.print("[yellow]1. Downloads folder[/yellow]")
                self.console.print("[yellow]2. Desktop[/yellow]")
                
                choice = input("\nEnter your choice (1-2): ")
                
                if choice == "1":
                    # Save to Downloads folder
                    save_path = str(Path.home() / "Downloads" / original_filename)
                elif choice == "2":
                    # Save to Desktop
                    save_path = str(Path.home() / "Desktop" / original_filename)
                else:
                    self.console.print("[red]Invalid choice. Saving to desktop.[/red]")
                    save_path = str(Path.home() / "Desktop" / original_filename)
        
                # Create directory if it doesn't exist
                Path(save_path).parent.mkdir(parents=True, exist_ok=True)
        
                # Save the document
                doc.save(save_path)
                self.console.print(f"\n[green]Feedback saved successfully to: {save_path}[/green]")
        
            except Exception as e:
                self.console.print(f"[red]Error saving feedback document: {str(e)}[/red]")
                # Final fallback - save to desktop
                desktop_path = Path.home() / 'Desktop' / f'resume_feedback_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                try:
                    doc.save(str(desktop_path))
                    self.console.print(f"\n[yellow]Feedback saved to desktop: {desktop_path}[/yellow]")
                except Exception as e2:
                    self.console.print(f"[red]Could not save feedback document: {str(e2)}[/red]")

    def download_resume(self, resume_url: str):
        """Download the resume from the provided URL."""
        try:
            import requests
            import tempfile
            import os

            if not resume_url:
                console.print("The Employer hasn't submitted a resume. This function will not work for this job.")
                return None

            # Create a temporary file to store the downloaded resume
            temp_dir = tempfile.gettempdir()
            temp_file = os.path.join(temp_dir, 'downloaded_resume.pdf')

            # Download the file
            response = requests.get(resume_url)
            response.raise_for_status()

            # Save the file
            with open(temp_file, 'wb') as f:
                f.write(response.content)

            return temp_file

        except requests.exceptions.RequestException as e:
            console.print(f"[red]Error downloading resume: {str(e)}[/red]")
            return None
        except Exception as e:
            console.print(f"[red]Error: {str(e)}[/red]")
            return None

def main():
    reviewer = ResumeReviewer()
    return reviewer

if __name__ == "__main__":
    main()